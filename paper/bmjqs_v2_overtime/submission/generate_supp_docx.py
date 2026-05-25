#!/usr/bin/env python3
"""Generate BMJ-formatted supplementary Word document with Figures S1, S2 and Tables S1, S2."""

import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGURES_DIR = '/home/user/ZOL/paper/bmjqs_v2_overtime/figures'

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Default font ---
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = Pt(20)
pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(text, bold=False, italic=False, size=12, space_after=6, space_before=0,
             alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def add_heading_caps(text, space_before=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p


def add_heading_lower(text, space_before=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_image(filename, width_inches=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(f'{FIGURES_DIR}/{filename}', width=Inches(width_inches))


def add_caption(label, text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing = Pt(16)
    rl = cap.add_run(label)
    rl.bold = True
    rl.font.name = 'Times New Roman'
    rl.font.size = Pt(11)
    rt = cap.add_run(' ' + text)
    rt.font.name = 'Times New Roman'
    rt.font.size = Pt(11)


def add_table(headers, rows, header_color='2171b5', bold_last_row=False):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color)
        tcPr.append(shd)

    for i, row in enumerate(rows, start=1):
        is_bold = bold_last_row and i == len(rows)
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(val))
            run.bold = is_bold
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)


# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = Pt(22)
run = p.add_run('Supplementary materials')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_para(
    'For: Where does operating-room overtime come from? A retrospective single-center study.',
    italic=True, size=11, space_after=12,
)
add_para('Tharwat H, Riebus M, [Ben], [Dieter], Martin N.', size=11, space_after=24)

# ============================================================
# TABLE S1 — CV by planned-duration bucket
# ============================================================
add_heading_lower('Supplementary Table S1. Coefficient of variation by planned-duration bucket')
add_table(
    headers=['Planned duration', 'n cases', 'CV (observed duration)', 'CV (planning deviation)'],
    rows=[
        ['<30 min',    '19,511', '0.61', '1.25'],
        ['31–60 min',  '28,674', '0.46', '1.07'],
        ['61–90 min',  '19,921', '0.36', '1.06'],
        ['91–180 min', '20,592', '0.35', '0.91'],
        ['>180 min',    '7,343', '0.42', '1.86'],
    ],
)
add_caption(
    'Table S1.',
    'Coefficient of variation (CV) of observed duration and of planning deviation, '
    'stratified by planned-duration bucket. Mid-length cases (61–180 min) carry the '
    'lowest CV. Very long cases (>180 min) carry the largest planning-deviation CV.',
)

# ============================================================
# TABLE S2 — Start-time deviation by room vs overtime
# ============================================================
add_heading_lower('Supplementary Table S2. Per-room start-time deviation vs overtime rate', space_before=18)
rooms_join = [
    ('OR11', 7482, 82.4, 319.1, 11.7),
    ('OR14', 6885, 78.7,  37.3,  3.5),
    ('OR01', 6577, 75.1,  38.0,  5.8),
    ('OR08', 1777, 71.9,  44.9, 13.6),
    ('OR02', 3502, 70.3,  42.2, 11.4),
    ('OR07', 4658, 70.0,  43.5,  6.9),
    ('OR05', 4886, 68.3,  54.3, 12.3),
    ('OR03', 4094, 65.9,  44.6,  9.6),
    ('OR09', 2637, 64.5,  48.9, 16.3),
    ('OR04', 4518, 63.5,  49.4, 10.7),
    ('OR15', 4323, 63.1,  40.5,  8.7),
    ('OR16', 4098, 62.6,  35.0,  8.9),
    ('OR18', 5293, 62.3,  36.9,  6.5),
    ('OR17', 5480, 61.8,  38.4,  6.8),
    ('OR06', 5217, 61.7,  42.1,  9.4),
    ('OR12', 2884, 54.5,  48.3, 13.3),
    ('OR13', 3298, 53.7,  48.6, 13.7),
    ('OR10', 1743, 46.1,  63.2, 32.9),
]
add_table(
    headers=['Room', 'n cases', 'Late starts (%)', 'Mean delay (min)', 'Overtime (%)'],
    rows=[[r, f'{n:,}', f'{l:.1f}', f'{d:.1f}', f'{o:.1f}'] for r, n, l, d, o in rooms_join],
)
add_caption(
    'Table S2.',
    'Per-room late-start rate and mean start delay (minutes) versus overtime rate. '
    'OR10 (lowest late-start rate, 46.1%) has the highest overtime (32.9%); '
    'OR14 (78.7% late, second-worst punctuality) has the lowest overtime (3.5%); '
    'OR11 (worst punctuality, 82.4% late, 319 min mean delay) sits at mid-pack overtime (11.7%).',
)

# ============================================================
# FIGURE S1 — Idle time
# ============================================================
add_heading_lower('Supplementary Figure S1. Idle time between consecutive cases', space_before=18)
add_image('figS1_idle_time.png', width_inches=6.0)
add_caption(
    'Figure S1.',
    'Distribution of inter-case idle time between consecutive cases in the same room. '
    'Mean 9.9 min, median 8 min, P95 25 min, P99 42 min. Turnover is fast and consistent.',
)

# ============================================================
# FIGURE S2 — End-time distribution
# ============================================================
add_heading_lower('Supplementary Figure S2. Timing of overtime case completions', space_before=18)
add_image('figS2_end_time_distribution.png', width_inches=6.0)
add_caption(
    'Figure S2.',
    'Number of overtime cases by end-time hour. The largest single block of completions '
    'falls in the 16:30–17:30 window, right after the day-shift handover. The distribution '
    'decays through the evening, with a thin tail past 22:00.',
)

# ============================================================
# SAVE
# ============================================================
output_path = '/home/user/ZOL/paper/bmjqs_v2_overtime/submission/manuscript_supplementary.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
