#!/usr/bin/env python3
"""Generate a BMJ Quality & Safety-formatted Word document from manuscript.md.

This parser reads the markdown manuscript and renders it with BMJ house style
(Times New Roman 12pt, double spacing, 2.5 cm margins, BOLD CAPS level-1
headings, bold lower-case level-2 headings). Parsing the markdown directly
keeps the .docx and the .md in lock-step so they cannot drift apart.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = '/home/user/ZOL/paper/bmjqs_v2_overtime/submission/manuscript.md'
OUT = '/home/user/ZOL/paper/bmjqs_v2_overtime/submission/manuscript.docx'

REF_RE = re.compile(r'\[\d+(?:[,–—-]\d+)*\]')
INLINE_RE = re.compile(r'(\*\*.*?\*\*|\*[^*]+?\*|\[\d+(?:[,–—-]\d+)*\])')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.line_spacing = Pt(24)


def add_runs(paragraph, text, size=12, base_bold=False, base_italic=False):
    """Render inline markdown (**bold**, *italic*, [ref] superscript)."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        elif REF_RE.fullmatch(part):
            run = paragraph.add_run(part[1:-1]); run.font.superscript = True
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)


def para(text, size=12, space_after=6, space_before=0, base_bold=False, base_italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = Pt(24)
    add_runs(p, text, size=size, base_bold=base_bold, base_italic=base_italic)
    return p


def heading(text, caps=False, size=12, space_before=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(24)
    run = p.add_run(text.upper() if caps else text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def is_table_sep(cells):
    return all(re.fullmatch(r':?-{2,}:?', c.strip()) for c in cells)


def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def add_table(rows):
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]  # rows[1] is the separator
    ncol = len(header)
    table = doc.add_table(rows=1 + len(body), cols=ncol)
    table.style = 'Table Grid'
    for j, cell_text in enumerate(header):
        cell = table.cell(0, j)
        cell.paragraphs[0].text = ''
        add_runs(cell.paragraphs[0], cell_text, size=11, base_bold=True)
    for i, row in enumerate(body, start=1):
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ''
            add_runs(cell.paragraphs[0], row[j] if j < len(row) else '', size=11)
    # tighten cell line spacing
    for r in table.rows:
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.line_spacing = Pt(13)
                p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.line_spacing = Pt(6)
    return table


# ---------------------------------------------------------------------------
# Parse the markdown
# ---------------------------------------------------------------------------
with open(SRC, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
first_h1_done = False
while i < len(lines):
    raw = lines[i].rstrip('\n')
    line = raw.strip()

    if not line or line == '---':
        i += 1
        continue

    # Tables: a block of consecutive lines starting with '|'
    if line.startswith('|'):
        block = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i].strip())
            i += 1
        if len(block) >= 2 and is_table_sep(split_row(block[1])):
            add_table(block)
        else:
            for b in block:
                para(b)
        continue

    if line.startswith('### '):
        heading(line[4:].strip(), caps=False, space_before=12)
        i += 1
        continue
    if line.startswith('## '):
        heading(line[3:].strip(), caps=True, space_before=16)
        i += 1
        continue
    if line.startswith('# '):
        text = line[2:].strip()
        if not first_h1_done:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(12)
            pf.line_spacing = Pt(24)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            first_h1_done = True
        else:
            heading(text, caps=True)
        i += 1
        continue

    # Bullet list item
    if line.startswith('- '):
        para('• ' + line[2:].strip(), space_after=3)
        i += 1
        continue

    # Numbered list item (RQs, references): keep "N. text"
    if re.match(r'^\d+\.\s', line):
        para(line, space_after=3)
        i += 1
        continue

    # Bracketed figure/table placeholders -> italic
    if line.startswith('[') and line.endswith(']'):
        para(line, base_italic=True, space_after=6, space_before=6)
        i += 1
        continue

    para(line)
    i += 1

doc.save(OUT)
print(f'Saved to {OUT}')
